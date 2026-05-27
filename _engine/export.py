from pathlib import Path
import re

def export_markdown(title, content):
    """导出 Markdown 格式"""
    stripped = content.strip()
    if stripped.startswith('# '):
        idx = stripped.find('\n')
        body = stripped[idx+1:].strip() if idx > 0 else ''
        return f"# {title}\n\n{body}"
    return f"# {title}\n\n{stripped}"

def export_html(title, content):
    """导出 HTML 格式"""
    md_content = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
body {{
    font-family: "Microsoft YaHei", "PingFang SC", sans-serif;
    max-width: 800px; margin: 40px auto; padding: 0 20px;
    line-height: 1.8; color: #333;
}}
h1 {{ border-bottom: 2px solid #6366f1; padding-bottom: 10px; }}
h2 {{ color: #4f46e5; }}
pre {{ background: #f5f5f5; padding: 16px; border-radius: 8px; overflow-x: auto; }}
code {{ background: #f0f0f0; padding: 2px 6px; border-radius: 4px; }}
blockquote {{ border-left: 4px solid #6366f1; margin: 0; padding-left: 16px; color: #666; }}
</style>
</head>
<body>
<h1>{title}</h1>
<div class="content">
{md_content.replace(chr(10), "<br>")}
</div>
</body>
</html>"""
    return html

def export_wechat_html(title, content, style="business"):
    """导出微信公众号 HTML（全部内联样式，5种风格）"""
    styles = {
        "business": {
            "bg": "#ffffff", "text": "#333333",
            "h1": "font-size:22px;font-weight:700;color:#1a1a1a;text-align:center;padding:24px 0 16px;border-bottom:1px solid #eee;margin-bottom:24px;",
            "h2": "font-size:17px;font-weight:600;color:#2d3436;padding:8px 0;margin:24px 0 12px;border-bottom:1px solid #dfe6e9;",
            "p": "font-size:15px;color:#333;line-height:1.85;margin:0 0 16px;letter-spacing:0.5px;",
            "quote": "font-size:14px;color:#636e72;background:#f5f6fa;padding:14px 18px;margin:20px 0;border-left:3px solid #0984e3;border-radius:0 6px 6px 0;line-height:1.75;",
            "code": "font-family:Menlo,Consolas,monospace;font-size:13px;background:#f0f0f0;padding:2px 6px;border-radius:3px;color:#e84393;",
            "list": "font-size:15px;color:#333;line-height:1.85;padding-left:20px;margin:0 0 16px;",
            "hr": "border:none;border-top:1px solid #eee;margin:24px 0;",
            "accent": "#0984e3"
        },
        "literary": {
            "bg": "#fefefe", "text": "#4a4a4a",
            "h1": "font-size:22px;font-weight:400;color:#2c3e50;text-align:center;padding:20px 0 12px;letter-spacing:2px;",
            "h2": "font-size:16px;font-weight:400;color:#8b7866;padding:10px 0;margin:20px 0 10px;text-align:center;",
            "p": "font-size:15px;color:#555;line-height:2;margin:0 0 20px;letter-spacing:0.8px;",
            "quote": "font-size:15px;color:#8b7866;text-align:center;padding:20px;margin:24px 0;font-style:italic;border-top:1px solid #e8d5c4;border-bottom:1px solid #e8d5c4;",
            "accent": "#c8956c"
        },
        "tech": {
            "bg": "#ffffff", "text": "#2d3436",
            "h1": "font-size:24px;font-weight:700;color:#0c0c0c;padding:24px 0 16px;line-height:1.3;",
            "h2": "font-size:18px;font-weight:600;color:#0984e3;padding:12px 0 8px;margin:28px 0 14px;border-left:4px solid #0984e3;padding-left:14px;",
            "p": "font-size:15px;color:#2d3436;line-height:1.8;margin:0 0 18px;",
            "quote": "font-size:14px;color:#636e72;background:#f8f9fa;padding:16px 20px;margin:20px 0;border-radius:8px;",
            "code": "font-family:Menlo,Consolas,monospace;font-size:13px;background:#2d3436;color:#55efc4;padding:3px 8px;border-radius:4px;",
            "accent": "#0984e3"
        },
        "youth": {
            "bg": "#fefefe", "text": "#2d3436",
            "h1": "font-size:23px;font-weight:700;color:#2d3436;padding:20px 0 10px;text-align:center;",
            "h2": "font-size:17px;font-weight:600;background:#fff5f5;padding:10px 14px;margin:24px 0 12px;border-radius:8px;color:#e74c3c;",
            "p": "font-size:15px;color:#444;line-height:1.9;margin:0 0 18px;",
            "quote": "font-size:15px;color:#e74c3c;text-align:center;padding:16px;margin:20px 0;background:#fff5f5;border-radius:12px;font-weight:500;",
            "code": "font-family:Menlo,monospace;font-size:13px;background:#ffeaa7;color:#d63031;padding:3px 8px;border-radius:4px;",
            "accent": "#ff6b6b"
        },
        "official": {
            "bg": "#ffffff", "text": "#333333",
            "h1": "font-size:24px;font-weight:700;color:#cc0000;text-align:center;padding:30px 0 20px;border-bottom:2px solid #cc0000;margin-bottom:28px;letter-spacing:1px;",
            "h2": "font-size:17px;font-weight:600;color:#cc0000;padding:6px 0;margin:24px 0 14px;border-bottom:1px solid #e0e0e0;",
            "p": "font-size:16px;color:#333;line-height:1.9;margin:0 0 18px;text-indent:2em;",
            "quote": "font-size:15px;color:#666;background:#f8f8f8;padding:16px 20px;margin:24px 0;border-left:4px solid #cc0000;line-height:1.8;",
            "code": "font-family:SimSun,serif;font-size:15px;color:#333;background:#f5f5f5;padding:2px 6px;border-radius:2px;",
            "accent": "#cc0000"
        }
    }
    cfg = styles.get(style, styles["business"])

    def format_wechat(text):
        lines = text.split('\n')
        result = []
        buf = []
        flush = lambda: result.append(''.join(buf)) if buf else None

        for line in lines:
            stripped = line.strip()
            if not stripped: flush(); buf = []; continue
            if stripped.startswith('### '):
                flush(); buf = []
                result.append('<h3 style="font-size:15px;font-weight:600;color:{0};padding:4px 0;margin:18px 0 8px;">{1}</h3>'.format(cfg['accent'], stripped[4:]))
                continue
            if stripped.startswith('## '):
                flush(); buf = []
                result.append('<h2 style="{0}">{1}</h2>'.format(cfg['h2'], stripped[3:]))
                continue
            if stripped.startswith('# '):
                flush(); buf = []
                result.append('<h1 style="{0}">{1}</h1>'.format(cfg['h1'], stripped[2:]))
                continue
            if stripped.startswith('> '):
                flush(); buf = []
                q_style = cfg.get('quote', cfg['p'])
                result.append('<blockquote style="{0}">{1}</blockquote>'.format(q_style, stripped[2:]))
                continue
            if stripped == '---' or stripped == '***':
                flush(); buf = []
                result.append('<hr style="{}">'.format(cfg.get('hr', 'border:none;border-top:1px solid #eee;margin:24px 0;')))
                continue
            if re.match(r'^\d+[\.\)] ', stripped):
                flush(); buf = []
                text = re.sub(r'^\d+[\.\)] ', '', stripped)
                result.append('<p style="{0}">{1}</p>'.format(cfg['p'], text))
                continue
            if stripped.startswith('- ') or stripped.startswith('* '):
                flush(); buf = []
                text = stripped[2:]
                result.append('<p style="{0}">- {1}</p>'.format(cfg['p'], text))
                continue
            buf.append(stripped)
        flush()
        return '\n'.join(result)

    body_html = format_wechat(content)
    html = f"""<section style="max-width:677px;margin:0 auto;padding:16px 12px 30px;background:{cfg['bg']};font-family:-apple-system,BlinkMacSystemFont,'PingFang SC','Microsoft YaHei',sans-serif;">
<h1 style="{cfg['h1']}">{title}</h1>
{body_html}
</section>"""
    return html

def save_file(path, content):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    return str(path)
def export_docx(title, content):
    from io import BytesIO
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm
    except ImportError:
        return None, 'python-docx 未安装，请运行 install_libs.bat'

    doc = Document()
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.5

    h = doc.add_heading(title, level=0)
    h.alignment = 1  # center

    for para in content.strip().split(chr(10)):
        line = para.strip()
        if not line: continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = (100,100,100) if hasattr(run.font.color,'rgb') else None
        else:
            doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue(), None

def export_pdf(title, content):
    from io import BytesIO
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return None, 'reportlab 未安装，请运行 install_libs.bat'

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()

    body_style = ParagraphStyle('CNBody', fontSize=11, leading=20, spaceAfter=12)
    h1_style = ParagraphStyle('CNH1', fontSize=20, leading=28, spaceAfter=16, textColor='#cc0000')
    h2_style = ParagraphStyle('CNH2', fontSize=15, leading=22, spaceAfter=10, spaceBefore=12)

    story = [Paragraph(title, h1_style), Spacer(1, 12)]

    for para in content.strip().split(chr(10)):
        line = para.strip()
        if not line: story.append(Spacer(1,6)); continue
        if line.startswith('# '): story.append(Paragraph(line[2:], h1_style))
        elif line.startswith('## '): story.append(Paragraph(line[3:], h2_style))
        elif line.startswith('### '): story.append(Paragraph('<b>' + line[4:] + '</b>', body_style))
        elif line.startswith('- '): story.append(Paragraph('• ' + line[2:], body_style))
        elif line.startswith('> '): story.append(Paragraph('<i>' + line[2:] + '</i>', body_style))
        else: story.append(Paragraph(line, body_style))

    doc.build(story)
    return buf.getvalue(), None